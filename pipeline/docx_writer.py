from docx import Document

def render_docx_from_state(state:dict, out_path:str="data/output/updated_resume.docx"):
    doc = Document()
    doc.add_heading("Updated Resume", level=0)
    doc.add_heading("Experience Bullets (updated)", level=1)
    for b in state["resume"]["bullets"]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_heading("ATS Summary", level=1)
    subs = state["features"]["subscores"]
    for k, v in subs.items():
        doc.add_paragraph(f"{k}: {v}%")
    doc.add_paragraph(f"Final ATS Score: {state['score']}%")
    doc.save(out_path)
